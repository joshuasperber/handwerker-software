#!/usr/bin/env python3
"""Convert Gesamtdokument Markdown (GFM tables) to a Word .docx with real tables."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")


def set_run_font(run, *, size=None, bold=None, italic=None, code=False):
    run.font.name = "Consolas" if code else "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas" if code else "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if code:
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_inline(paragraph, text, base_size=11):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=base_size)
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, code=True)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=base_size, italic=True)
        else:
            label = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).group(1)
            run = paragraph.add_run(label)
            set_run_font(run, size=base_size)
            run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
            run.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size)


def shade_cell(cell, color_hex="E8EEF5"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tc_pr.append(borders)


def parse_table_row(line: str) -> list[str]:
    raw = line.strip()
    if raw.startswith("|"):
        raw = raw[1:]
    if raw.endswith("|"):
        raw = raw[:-1]
    return [c.strip() for c in raw.split("|")]


def is_sep_row(cells: list[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells)


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def add_table(rows: list[list[str]]):
        if not rows:
            return
        cols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(cols):
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                add_inline(p, row[c_idx] if c_idx < len(row) else "", base_size=9.5)
                set_cell_borders(cell)
                if r_idx == 0:
                    shade_cell(cell)
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(buf))
            set_run_font(run, code=True)
            continue

        if line.strip().startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = parse_table_row(lines[i])
                if not is_sep_row(cells):
                    rows.append(cells)
                i += 1
            add_table(rows)
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            p = doc.add_heading(level=level)
            p.clear()
            sizes = {1: 20, 2: 16, 3: 13, 4: 12}
            add_inline(p, m.group(2), base_size=sizes[level])
            for run in p.runs:
                run.bold = True
            i += 1
            continue

        if line.strip().startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, lines[i].strip()[2:])
                i += 1
            continue

        if re.match(r"^\d+\.\s+", line.strip()):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                p = doc.add_paragraph(style="List Number")
                add_inline(p, re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            continue

        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    doc.save(out_path)


def main() -> int:
    root = Path(__file__).resolve().parent
    md = root / "gesamtdokument-jomaster.md"
    out = root / "gesamtdokument-jomaster.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    convert(md, out)
    print(f"Wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
